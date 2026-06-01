from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"D:\xwechat_files\wxid_kmb6d6zs5vhi21_55b1\msg\file\2026-03\cosmic拆分\diff_service_agent_simplified_300.docx")


SECTIONS = [
    ("基础数据接入", "华为智能板5G小区清单接入", "按日接入华为智能板小区清单，识别小区标识、区域、站点和基础属性。"),
    ("基础数据接入", "华为全网5G小区级数据接入", "接入华为全网小区级性能数据，保留覆盖、容量、接入和保持相关指标。"),
    ("基础数据接入", "华为全网5G 5QI级数据接入", "接入华为分5QI业务体验数据，支撑按业务等级识别体验差异。"),
    ("基础数据接入", "华为全网5G DU5QI级数据接入", "接入华为DU侧分5QI数据，补充用户面和承载侧质量分析。"),
    ("基础数据接入", "华为智能板5G PRS级数据接入", "接入华为PRS级业务指标，用于重点业务感知和异常识别。"),
    ("基础数据接入", "华为智能板5G管B级数据接入", "接入华为管B侧指标，关联网络资源与业务体验变化。"),
    ("基础数据接入", "中兴全网5G小区清单接入", "按日接入中兴小区清单，统一小区、区域、厂家和制式属性。"),
    ("基础数据接入", "中兴全网5G DRA参数数据接入", "接入中兴DRA参数数据，保留关键策略、切片和业务承载参数。"),
    ("基础数据接入", "中兴全网5G分5QI KPI指标接入", "接入中兴分5QI KPI指标，覆盖接入、时延、速率和质量类指标。"),
    ("基础数据接入", "中兴全网5G分5QI接通掉线指标接入", "接入中兴分5QI接通与掉线数据，支撑连接稳定性分析。"),
    ("基础数据接入", "中兴智能板5G分业务指标接入", "接入中兴智能板分业务指标，识别视频、直播、游戏等业务体验特征。"),
    ("基础数据接入", "中兴智能板5G 5QI6分业务指标接入", "接入中兴5QI6重点业务指标，支撑保障类业务专项分析。"),
    ("指标数据建模", "5G全量工参数据建模", "汇聚厂家小区清单与工参，形成统一5G小区基础模型。"),
    ("指标数据建模", "5G全量小区分5QI计算数据建模", "按小区和5QI维度计算体验指标，生成可分析的明细数据。"),
    ("指标数据建模", "5G全量小区分5QI统计数据建模", "按区域、厂家、时间和5QI维度汇总指标，形成统计结果。"),
    ("指标数据建模", "5G智能板分业务类型计算数据建模", "按业务类型计算小区体验指标，保留业务量、速率、时延和质量特征。"),
    ("指标数据建模", "5G智能板分业务类型统计数据建模", "按业务类型汇总统计，支撑专题看板和趋势分析。"),
    ("厂家接口对接", "华为厂家放号查询接口对接", "对接华为放号查询接口，获取小区放号状态、策略和预测结果。"),
    ("厂家接口对接", "中兴厂家放号查询接口对接", "对接中兴放号查询接口，获取小区放号状态、策略和预测结果。"),
    ("前端模块呈现", "分5QI小区级指标前台呈现", "在前台展示小区级分5QI指标，支持筛选、排序、钻取和导出。"),
    ("前端模块呈现", "分5QI统计级指标前台呈现", "在前台展示分5QI统计结果，支持区域、厂家和时间维度分析。"),
    ("前端模块呈现", "分业务类型小区级指标前台呈现", "在前台展示小区级分业务体验指标，支持重点业务识别。"),
    ("前端模块呈现", "分业务类型统计级指标前台呈现", "在前台展示业务类型统计结果，支持趋势、排名和占比分析。"),
    ("前端模块呈现", "放号预测配置前台呈现", "提供放号预测参数配置，支持区域、厂家、阈值和策略维护。"),
    ("前端模块呈现", "放号预测小区级结果前台呈现", "展示小区级放号预测结果，支持状态核验和结果导出。"),
    ("前端模块呈现", "放号预测场景级结果前台呈现", "展示场景级放号预测结果，支持场景聚合和风险识别。"),
    ("前端模块呈现", "质差小区前台呈现", "展示质差小区清单，支持原因分类、指标明细和处置跟踪。"),
    ("前端模块呈现", "质差小区推送工单", "将质差小区按规则推送至工单系统，支撑闭环处理。"),
    ("前端模块呈现", "疑似直播聚集小区目标与剔除清单呈现", "展示疑似直播聚集小区目标清单和剔除清单，支持人工核验。"),
    ("首页与问答库", "首页看板与问答库综合支撑", "提供首页总览、趋势提示、口径说明和问答库查询。"),
]


BASE_OPS = [
    ("接收", "支持文件或接口方式获取数据，记录批次、来源、时间和处理状态。"),
    ("校验", "校验必填字段、时间范围、厂家标识和关键指标完整性。"),
    ("解析", "解析原始字段并提取小区、区域、业务类型和指标值。"),
    ("映射", "将厂家字段映射为统一字段口径，便于跨厂家对比。"),
    ("存储", "按日期、厂家、区域和业务维度保存明细数据。"),
    ("更新", "支持同批次覆盖更新，避免重复导入造成指标偏差。"),
    ("查询", "按时间、区域、厂家、小区和业务维度查询数据。"),
    ("统计", "计算总量、均值、占比、排名和异常数量等统计结果。"),
    ("展示", "以列表、趋势、排名和详情方式展示核心指标。"),
    ("导出", "支持导出当前筛选结果，保留统计口径和生成时间。"),
]


OVERRIDES = {
    "华为厂家放号查询接口对接": [
        ("配置", "维护华为接口地址、鉴权参数、超时阈值和调用开关。"),
        ("鉴权", "按接口规范生成请求凭证，记录鉴权成功或失败原因。"),
        ("请求", "按区域、厂家和小区范围发起放号状态查询请求。"),
        ("重试", "对超时或临时失败请求执行有限重试，避免重复扩大调用量。"),
        ("解析", "解析放号状态、预测等级、策略建议和返回时间。"),
        ("关联", "将接口结果关联到统一小区模型和前台展示对象。"),
        ("缓存", "缓存近期查询结果，降低重复查询对厂家接口的压力。"),
        ("告警", "对接口异常、返回缺字段和状态不一致生成提示。"),
        ("展示", "在放号预测页面展示华为小区结果和关键解释字段。"),
        ("留痕", "记录请求参数、响应摘要、操作人和调用时间，便于追溯。"),
    ],
    "中兴厂家放号查询接口对接": [
        ("配置", "维护中兴接口地址、鉴权参数、超时阈值和调用开关。"),
        ("鉴权", "按接口规范生成请求凭证，记录鉴权成功或失败原因。"),
        ("请求", "按区域、厂家和小区范围发起放号状态查询请求。"),
        ("重试", "对超时或临时失败请求执行有限重试，避免重复扩大调用量。"),
        ("解析", "解析放号状态、预测等级、策略建议和返回时间。"),
        ("关联", "将接口结果关联到统一小区模型和前台展示对象。"),
        ("缓存", "缓存近期查询结果，降低重复查询对厂家接口的压力。"),
        ("告警", "对接口异常、返回缺字段和状态不一致生成提示。"),
        ("展示", "在放号预测页面展示中兴小区结果和关键解释字段。"),
        ("留痕", "记录请求参数、响应摘要、操作人和调用时间，便于追溯。"),
    ],
    "质差小区推送工单": [
        ("识别", "按质差规则识别待派单小区，保留触发指标和触发时间。"),
        ("筛选", "支持按区域、厂家、业务类型和质差等级筛选待推送清单。"),
        ("合并", "对同小区同问题的重复记录进行合并，减少重复派单。"),
        ("补全", "补全小区名称、站点、区域、责任单位和关键指标信息。"),
        ("预览", "推送前预览工单标题、问题描述、指标证据和处理建议。"),
        ("推送", "将确认后的质差小区推送至工单系统并接收工单编号。"),
        ("回写", "将工单编号、推送状态和失败原因回写到分析结果。"),
        ("跟踪", "跟踪工单处理状态，支持未处理、处理中、已闭环查询。"),
        ("撤销", "对误推送或已恢复小区支持撤销或关闭推送标记。"),
        ("统计", "统计推送数量、成功率、闭环率和超时未处理数量。"),
    ],
}


HIGHLIGHTS = {
    "基础数据接入": "输入、校验、字段映射、入库和基础查询。",
    "指标数据建模": "统一口径、指标计算、维度汇总和结果沉淀。",
    "厂家接口对接": "接口配置、调用、解析、缓存和异常留痕。",
    "前端模块呈现": "筛选、列表、趋势、详情、导出和闭环操作。",
    "首页与问答库": "总览、指标解释、常见问题和辅助查询。",
}


def apply_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    for style_name, size in [("Normal", 9.5), ("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 10.5)]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_run_font(p.add_run("差异化服务智能体需求精简版"), 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_run_font(p.add_run("用于 COSMIC 功能过程拆分：控制为 300 条候选功能过程"), 10, False, (91, 105, 135))

    doc.add_heading("使用说明", level=1)
    for text in [
        "本版本从原始业务需求中抽取核心功能特点，删除明细字段字典、历史变更说明、重复表格和样例数据。",
        "全文按 30 个业务主题组织，每个主题保留 10 条功能过程候选项，共 300 条。",
        "每条以“功能过程 FP-xxx”开头，便于拆分系统稳定识别；如需减少数量，可按业务主题整组删除。",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("范围概览", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, text in enumerate(["业务域", "主题数量", "保留功能特点"]):
        cell = table.rows[0].cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            apply_run_font(run, 9, True)

    for domain in ["基础数据接入", "指标数据建模", "厂家接口对接", "前端模块呈现", "首页与问答库"]:
        cells = table.add_row().cells
        cells[0].text = domain
        cells[1].text = str(sum(1 for d, _, _ in SECTIONS if d == domain))
        cells[2].text = HIGHLIGHTS[domain]

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    apply_run_font(run, 9)

    items = 0
    current_domain = None
    doc.add_heading("功能过程候选清单（共300条）", level=1)
    for domain, topic, desc in SECTIONS:
        if current_domain != domain:
            current_domain = domain
            doc.add_heading(domain, level=2)
        doc.add_heading(topic, level=3)
        doc.add_paragraph("主题特点：" + desc)
        for verb, detail in OVERRIDES.get(topic, BASE_OPS):
            items += 1
            p = doc.add_paragraph()
            apply_run_font(p.add_run(f"功能过程 FP-{items:03d}：{verb}{topic}。"), 9.5, True)
            apply_run_font(p.add_run("功能特点：" + detail), 9.5)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.08

    doc.core_properties.title = "差异化服务智能体需求精简版"
    doc.core_properties.subject = "COSMIC 功能过程拆分 300 条候选项"
    doc.core_properties.comments = "Generated concise version; original document unchanged."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    with ZipFile(OUT) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    assert "差异化服务智能体需求精简版" in xml
    assert "????" not in xml
    print(f"created={OUT}")
    print(f"items={items}")
    print("chinese_ok=True")


if __name__ == "__main__":
    main()
